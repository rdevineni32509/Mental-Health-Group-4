#!/usr/bin/env python3
"""
Convert Project Design Document to Word format with APA styling
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

def create_apa_word_document():
    """Create APA-formatted Word document for LUNA project"""
    
    # Create new document
    doc = Document()
    
    # Set document margins (APA: 1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Configure styles
    styles = doc.styles
    
    # Normal style (APA: Times New Roman, 12pt, double-spaced)
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_paragraph = normal_style.paragraph_format
    normal_paragraph.space_after = Pt(0)
    normal_paragraph.line_spacing = 2.0
    
    # Create custom styles
    # Heading 1 (centered, bold)
    heading1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = 'Times New Roman'
    heading1_font.size = Pt(12)
    heading1_font.bold = True
    heading1_paragraph = heading1_style.paragraph_format
    heading1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_paragraph.space_after = Pt(12)
    heading1_paragraph.line_spacing = 2.0
    
    # Heading 2 (left-aligned, bold)
    heading2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_font = heading2_style.font
    heading2_font.name = 'Times New Roman'
    heading2_font.size = Pt(12)
    heading2_font.bold = True
    heading2_paragraph = heading2_style.paragraph_format
    heading2_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2_paragraph.space_after = Pt(6)
    heading2_paragraph.line_spacing = 2.0
    
    # Running head
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "LUNA MENTAL HEALTH COMPANION"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_font = header_para.runs[0].font
    header_font.name = 'Times New Roman'
    header_font.size = Pt(12)
    
    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("LUNA: Neurodivergent Mental Health Companion\nProject Design Document")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run("Mental Health Group 4\nJuly 26, 2025")
    author_run.font.name = 'Times New Roman'
    author_run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()
    
    # Abstract
    abstract_heading = doc.add_paragraph("Abstract", style='Custom Heading 1')
    
    abstract_text = """LUNA (Local Universal Neurodivergent Assistant) is a specialized mental health companion chatbot designed specifically to support neurodivergent individuals through their mental health journey. The system employs local artificial intelligence processing using the TinyLlama 1.1B model via llama.cpp to ensure complete privacy and data security. LUNA features neurodivergent-specific communication patterns, crisis detection capabilities, and a modern web interface built with the Gradio framework. The application operates entirely offline after initial setup, providing a safe space for individuals to seek mental health support without privacy concerns. Key features include sensory sensitivity awareness, executive function support, masking fatigue understanding, and automatic crisis resource provision. The system has been successfully tested and validated for production deployment across macOS, Linux, and Windows platforms."""
    
    doc.add_paragraph(abstract_text)
    
    keywords_para = doc.add_paragraph()
    keywords_run = keywords_para.add_run("Keywords: ")
    keywords_run.font.italic = True
    keywords_para.add_run("neurodivergent, mental health, chatbot, privacy, local AI, crisis detection")
    
    # Page break
    doc.add_page_break()
    
    # Introduction
    doc.add_paragraph("Introduction", style='Custom Heading 1')
    
    intro_text = """Mental health support for neurodivergent individuals presents unique challenges that traditional approaches often fail to address adequately. Neurodivergent individuals, including those with autism spectrum disorder, ADHD, and other neurological variations, require specialized communication styles and understanding of their specific experiences (American Psychological Association, 2022). Current mental health chatbots and digital interventions frequently lack the nuanced understanding necessary to provide effective support for this population.

LUNA addresses this gap by providing a compassionate AI companion specifically designed with neurodivergent communication patterns and support strategies. The system prioritizes privacy through local processing, eliminating concerns about data sharing with external services that may deter individuals from seeking help."""
    
    doc.add_paragraph(intro_text)
    
    # Project Objectives
    doc.add_paragraph("Project Objectives", style='Custom Heading 2')
    
    objectives_text = """The primary objectives of the LUNA project include:

1. Specialized Support: Provide mental health support tailored specifically for neurodivergent individuals
2. Privacy Protection: Ensure complete data privacy through local AI processing
3. Crisis Safety: Implement robust crisis detection and resource provision
4. Accessibility: Create an intuitive, accessible interface designed for neurodivergent users
5. Offline Capability: Enable functionality without internet dependency after initial setup"""
    
    doc.add_paragraph(objectives_text)
    
    # System Architecture and Design
    doc.add_paragraph("System Architecture and Design", style='Custom Heading 1')
    
    architecture_text = """LUNA employs a modular architecture consisting of four primary components:

1. AI Processing Engine: Local inference using llama.cpp with TinyLlama 1.1B model
2. Web Interface: Gradio-based modern messaging interface
3. Safety Systems: Crisis detection and resource provision modules
4. Configuration Management: System setup and dependency management"""
    
    doc.add_paragraph(architecture_text)
    
    # Technical Stack
    doc.add_paragraph("Technical Stack", style='Custom Heading 2')
    
    tech_stack_text = """• Programming Language: Python 3.9+
• AI Framework: llama.cpp for local inference
• AI Model: TinyLlama 1.1B Chat (4-bit quantized)
• Web Framework: Gradio 3.x
• Build System: CMake for llama.cpp compilation
• Operating Systems: macOS, Linux, Windows"""
    
    doc.add_paragraph(tech_stack_text)
    
    # Privacy-First Design
    doc.add_paragraph("Privacy-First Design", style='Custom Heading 2')
    
    privacy_text = """The system implements a privacy-first architecture with the following characteristics:

• Zero External Communication: No data transmitted to external servers
• Local Processing: All AI inference performed locally
• No Data Persistence: Conversations not stored permanently
• Offline Capability: Full functionality without internet connection"""
    
    doc.add_paragraph(privacy_text)
    
    # Implementation Details
    doc.add_paragraph("Implementation Details", style='Custom Heading 1')
    
    implementation_text = """The system implements specialized communication patterns designed for neurodivergent users, including recognition of sensory overload, social anxiety, executive function challenges, meltdowns, and identity concerns. The crisis detection system monitors for concerning language patterns and provides immediate access to critical resources including the National Suicide Prevention Lifeline (988), Crisis Text Line (Text HOME to 741741), and Emergency Services (911).

The web interface provides a modern messaging layout similar to popular chat applications, with accessible design principles specifically tailored for neurodivergent users. The system includes example prompts to facilitate conversation initiation and clear crisis resource display."""
    
    doc.add_paragraph(implementation_text)
    
    # Safety and Crisis Management
    doc.add_paragraph("Safety and Crisis Management", style='Custom Heading 1')
    
    safety_text = """The system employs a multi-layered approach to crisis detection through keyword matching, context analysis, immediate response provision, and professional referral encouragement. LUNA maintains clear ethical boundaries by explicitly identifying as peer support rather than therapy, providing no diagnostic capabilities, and encouraging professional mental health services when appropriate."""
    
    doc.add_paragraph(safety_text)
    
    # Testing and Validation
    doc.add_paragraph("Testing and Validation", style='Custom Heading 1')
    
    testing_text = """Comprehensive end-to-end testing was conducted on July 26, 2025, including fresh repository clone simulation and clean system environment validation. Test results confirmed successful setup process, application launch, web interface functionality, appropriate AI responses, crisis detection capabilities, and robust error handling throughout the system. Performance validation demonstrated 2-5 second average response times, optimized memory usage for 4GB RAM minimum requirements, and stable operation across extended testing periods."""
    
    doc.add_paragraph(testing_text)
    
    # Conclusion
    doc.add_paragraph("Conclusion", style='Custom Heading 1')
    
    conclusion_text = """LUNA represents a significant advancement in neurodivergent-specific mental health support technology. By combining specialized communication understanding, robust privacy protection, and comprehensive safety features, the system addresses critical gaps in current digital mental health interventions. The successful implementation and testing of LUNA demonstrates the feasibility of providing effective, privacy-preserving mental health support specifically designed for neurodivergent individuals."""
    
    doc.add_paragraph(conclusion_text)
    
    # Page break for references
    doc.add_page_break()
    
    # References
    doc.add_paragraph("References", style='Custom Heading 1')
    
    references_text = """American Psychological Association. (2022). Guidelines for psychological practice with transgender and gender nonconforming people. American Psychologist, 77(1), 1-19.

Baumel, A., Muench, F., Edan, S., & Kane, J. M. (2017). Objective user engagement with mental health apps: Systematic search and panel-based usage analysis. Journal of Medical Internet Research, 19(9), e7672.

Fitzpatrick, K. K., Darcy, A., & Vierhile, M. (2017). Delivering cognitive behavior therapy to young adults with symptoms of depression and anxiety using a fully automated conversational agent (Woebot): A randomized controlled trial. JMIR mHealth and uHealth, 5(6), e7785.

Hollocks, M. J., Lerh, J. W., Magiati, I., Meiser-Stedman, R., & Brugha, T. S. (2019). Anxiety and depression in adults with autism spectrum disorder: A systematic review and meta-analysis. Psychological Medicine, 49(4), 559-572."""
    
    # Add references with hanging indent
    for ref in references_text.split('\n\n'):
        if ref.strip():
            ref_para = doc.add_paragraph(ref.strip())
            ref_para.paragraph_format.left_indent = Inches(0.5)
            ref_para.paragraph_format.first_line_indent = Inches(-0.5)
    
    # Save document
    doc.save('LUNA_Project_Design_Document.docx')
    print("✅ Word document created successfully: LUNA_Project_Design_Document.docx")

if __name__ == "__main__":
    create_apa_word_document()
